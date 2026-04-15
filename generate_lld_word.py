"""
Converts BlueLine_LLD.md to a professionally formatted Word document.
Run: python generate_lld_word.py
Output: BlueLine_LLD.docx
"""

import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

INPUT_FILE  = Path(__file__).parent / "BlueLine_LLD.md"
OUTPUT_FILE = Path(__file__).parent / "BlueLine_LLD.docx"

# ── Colour palette ────────────────────────────────────────────
BLUE_DARK   = RGBColor(0x1A, 0x3A, 0x6B)   # title / H1
BLUE_MID    = RGBColor(0x1E, 0x5F, 0xAF)   # H2
BLUE_LIGHT  = RGBColor(0x2E, 0x86, 0xC1)   # H3
GREY_BG     = RGBColor(0xF2, 0xF2, 0xF2)   # table header bg
CODE_BG     = RGBColor(0xF6, 0xF8, 0xFA)   # code block bg
CODE_FG     = RGBColor(0x24, 0x29, 0x2E)   # code text
TABLE_LINE  = RGBColor(0xBD, 0xC3, 0xC7)   # table border


# ── Helper: set paragraph shading ────────────────────────────
def set_paragraph_shading(paragraph, fill_colour_hex: str):
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_colour_hex)
    pPr.append(shd)


def set_cell_shading(cell, fill_colour_hex: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_colour_hex)
    tcPr.append(shd)


def set_cell_border(cell, **kwargs):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        tag = OxmlElement(f"w:{ edge }")
        tag.set(qn("w:val"),   kwargs.get("val",   "single"))
        tag.set(qn("w:sz"),    kwargs.get("sz",    "4"))
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), kwargs.get("color", "BDBDBD"))
        tcBorders.append(tag)
    tcPr.append(tcBorders)


def add_page_number(doc: Document):
    """Adds 'Page X of Y' footer to every page."""
    section = doc.sections[0]
    footer  = section.footer
    para    = footer.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    para.clear()
    run = para.add_run("Page ")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    # PAGE field
    fldChar = OxmlElement("w:fldChar")
    fldChar.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.text = "PAGE"
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    run2 = para.add_run()
    run2.font.size = Pt(9)
    run2.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    run2._r.append(fldChar)
    run2._r.append(instrText)
    run2._r.append(fldChar2)

    run3 = para.add_run(" of ")
    run3.font.size = Pt(9)
    run3.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    # NUMPAGES field
    fldChar3 = OxmlElement("w:fldChar")
    fldChar3.set(qn("w:fldCharType"), "begin")
    instrText2 = OxmlElement("w:instrText")
    instrText2.text = "NUMPAGES"
    fldChar4 = OxmlElement("w:fldChar")
    fldChar4.set(qn("w:fldCharType"), "end")
    run4 = para.add_run()
    run4.font.size = Pt(9)
    run4.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    run4._r.append(fldChar3)
    run4._r.append(instrText2)
    run4._r.append(fldChar4)


# ── Title page ────────────────────────────────────────────────
def add_title_page(doc: Document):
    # Top spacer
    for _ in range(6):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("PROJECT BLUELINE")
    run.font.size    = Pt(32)
    run.font.bold    = True
    run.font.color.rgb = BLUE_DARK

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = subtitle.add_run("Low Level Design Document")
    run2.font.size = Pt(20)
    run2.font.color.rgb = BLUE_MID

    doc.add_paragraph()

    line = doc.add_paragraph()
    line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = line.add_run("─" * 40)
    run3.font.color.rgb = BLUE_LIGHT

    doc.add_paragraph()

    for label, value in [
        ("Version",    "1.1"),
        ("Date",       "2026-04-15"),
        ("Status",     "Draft for Review"),
        ("Prepared by","Project BlueLine Team"),
    ]:
        meta = doc.add_paragraph()
        meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = meta.add_run(f"{label}: ")
        r1.font.bold      = True
        r1.font.size      = Pt(11)
        r1.font.color.rgb = BLUE_DARK
        r2 = meta.add_run(value)
        r2.font.size      = Pt(11)
        r2.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    doc.add_page_break()


# ── Markdown → Word parser ────────────────────────────────────
def apply_inline(run_parent, text: str):
    """
    Handles inline **bold**, `code`, and plain text in a paragraph.
    run_parent is a paragraph object.
    """
    # Split on **bold** and `code` markers
    parts = re.split(r'(\*\*[^*]+\*\*|`[^`]+`)', text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            r = run_parent.add_run(part[2:-2])
            r.bold = True
        elif part.startswith("`") and part.endswith("`"):
            r = run_parent.add_run(part[1:-1])
            r.font.name  = "Consolas"
            r.font.size  = Pt(9)
            r.font.color.rgb = RGBColor(0xC7, 0x25, 0x4F)
        else:
            if part:
                run_parent.add_run(part)


def parse_table(doc: Document, table_lines: list[str]):
    """Renders a markdown table as a formatted Word table."""
    rows = []
    for line in table_lines:
        if re.match(r'^\s*\|[-| :]+\|\s*$', line):
            continue   # skip separator row
        cells = [c.strip() for c in line.strip().strip("|").split("|")]
        if cells:
            rows.append(cells)

    if not rows:
        return

    col_count = max(len(r) for r in rows)
    tbl = doc.add_table(rows=len(rows), cols=col_count)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    for i, row_data in enumerate(rows):
        for j, cell_text in enumerate(row_data):
            if j >= col_count:
                break
            cell = tbl.cell(i, j)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            para = cell.paragraphs[0]
            para.clear()

            # Clean markdown links
            cell_text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', cell_text)

            if i == 0:
                # Header row
                set_cell_shading(cell, "1A3A6B")
                r = para.add_run(cell_text)
                r.bold            = True
                r.font.color.rgb  = RGBColor(0xFF, 0xFF, 0xFF)
                r.font.size       = Pt(9)
            else:
                set_cell_shading(cell, "FFFFFF" if i % 2 == 0 else "F5F8FC")
                apply_inline(para, cell_text)
                for run in para.runs:
                    run.font.size = Pt(9)

            set_cell_border(cell, color="C8D4E3")

    doc.add_paragraph()


def parse_code_block(doc: Document, lines: list[str], lang: str):
    """Renders a fenced code block with grey background."""
    # Opening label
    if lang and lang.strip():
        lbl = doc.add_paragraph()
        lbl.paragraph_format.space_before = Pt(4)
        lbl.paragraph_format.space_after  = Pt(0)
        r = lbl.add_run(f"  {lang.strip()}")
        r.font.size       = Pt(8)
        r.font.bold       = True
        r.font.color.rgb  = RGBColor(0x55, 0x55, 0x55)
        set_paragraph_shading(lbl, "DEE3EA")

    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        p.paragraph_format.left_indent  = Cm(0.5)
        set_paragraph_shading(p, "F6F8FA")
        r = p.add_run(line.rstrip())
        r.font.name  = "Consolas"
        r.font.size  = Pt(8)
        r.font.color.rgb = CODE_FG

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(0)
    spacer.paragraph_format.space_after  = Pt(6)
    set_paragraph_shading(spacer, "F6F8FA")


def add_horizontal_rule(doc: Document):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"),   "single")
    bottom.set(qn("w:sz"),    "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "2E86C1")
    pBdr.append(bottom)
    pPr.append(pBdr)


def convert_md_to_docx(md_path: Path, out_path: Path):
    doc = Document()

    # ── Page setup ──
    section = doc.sections[0]
    section.page_width      = Inches(8.5)
    section.page_height     = Inches(11)
    section.left_margin     = Inches(1.0)
    section.right_margin    = Inches(1.0)
    section.top_margin      = Inches(1.0)
    section.bottom_margin   = Inches(1.0)

    # ── Default body font ──
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    # ── Footer ──
    add_page_number(doc)

    # ── Title page ──
    add_title_page(doc)

    lines = md_path.read_text(encoding="utf-8").splitlines()

    i = 0
    table_lines: list[str] = []
    code_lines:  list[str] = []
    in_code  = False
    code_lang = ""
    in_table = False

    while i < len(lines):
        line = lines[i]

        # ── Code block ──────────────────────────────────────────
        if line.startswith("```"):
            if not in_code:
                in_code   = True
                code_lang = line[3:].strip()
                code_lines = []
            else:
                parse_code_block(doc, code_lines, code_lang)
                in_code    = False
                code_lang  = ""
                code_lines = []
            i += 1
            continue

        if in_code:
            code_lines.append(line)
            i += 1
            continue

        # ── Table detection ─────────────────────────────────────
        is_table_line = line.strip().startswith("|") and "|" in line[1:]

        if is_table_line:
            in_table = True
            table_lines.append(line)
            i += 1
            continue
        elif in_table:
            parse_table(doc, table_lines)
            table_lines = []
            in_table    = False
            # fall through to process current line

        # ── Horizontal rule ─────────────────────────────────────
        if re.match(r'^-{3,}$', line.strip()):
            add_horizontal_rule(doc)
            i += 1
            continue

        # ── Headings ─────────────────────────────────────────────
        if line.startswith("# ") and not line.startswith("## "):
            # Skip the document title line (already on title page)
            if "Project BlueLine" in line and "LLD" in line:
                i += 1
                continue
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(18)
            p.paragraph_format.space_after  = Pt(6)
            r = p.add_run(line[2:].strip())
            r.font.size    = Pt(20)
            r.font.bold    = True
            r.font.color.rgb = BLUE_DARK
            i += 1
            continue

        if line.startswith("## "):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after  = Pt(4)
            r = p.add_run(line[3:].strip())
            r.font.size    = Pt(15)
            r.font.bold    = True
            r.font.color.rgb = BLUE_MID
            # Blue bottom border
            pPr   = p._p.get_or_add_pPr()
            pBdr  = OxmlElement("w:pBdr")
            btm   = OxmlElement("w:bottom")
            btm.set(qn("w:val"),   "single")
            btm.set(qn("w:sz"),    "4")
            btm.set(qn("w:space"), "1")
            btm.set(qn("w:color"), "2E86C1")
            pBdr.append(btm)
            pPr.append(pBdr)
            i += 1
            continue

        if line.startswith("### "):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after  = Pt(3)
            r = p.add_run(line[4:].strip())
            r.font.size    = Pt(12)
            r.font.bold    = True
            r.font.color.rgb = BLUE_LIGHT
            i += 1
            continue

        if line.startswith("#### "):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after  = Pt(2)
            r = p.add_run(line[5:].strip())
            r.font.size    = Pt(11)
            r.font.bold    = True
            r.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
            i += 1
            continue

        # ── Bullet lists ─────────────────────────────────────────
        if re.match(r'^(\s*)[-*] ', line):
            indent = len(line) - len(line.lstrip())
            text   = re.sub(r'^(\s*)[-*] ', '', line)
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.left_indent  = Cm(0.5 + indent * 0.2)
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after  = Pt(1)
            p.clear()
            bullet_run = p.add_run("• ")
            bullet_run.font.color.rgb = BLUE_MID
            bullet_run.font.bold      = True
            apply_inline(p, text.strip())
            for r in p.runs[1:]:
                r.font.size = Pt(10)
            i += 1
            continue

        # ── Numbered lists ───────────────────────────────────────
        if re.match(r'^\d+\. ', line):
            text = re.sub(r'^\d+\. ', '', line)
            num  = re.match(r'^(\d+)\.', line).group(1)
            p = doc.add_paragraph()
            p.paragraph_format.left_indent  = Cm(0.5)
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after  = Pt(1)
            nr = p.add_run(f"{num}. ")
            nr.font.bold      = True
            nr.font.color.rgb = BLUE_MID
            apply_inline(p, text.strip())
            i += 1
            continue

        # ── Blank line ───────────────────────────────────────────
        if not line.strip():
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after  = Pt(3)
            i += 1
            continue

        # ── Metadata lines (Version:, Date:, etc.) ───────────────
        if re.match(r'^\*\*(Version|Date|Status|Prepared by)\*\*', line):
            i += 1
            continue

        # ── Block quote / note lines ─────────────────────────────
        if line.startswith("> "):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent  = Cm(1.0)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after  = Pt(4)
            set_paragraph_shading(p, "EBF5FB")
            r = p.add_run("  ")
            apply_inline(p, line[2:].strip())
            for r in p.runs:
                r.font.size       = Pt(9.5)
                r.font.color.rgb  = RGBColor(0x1A, 0x5C, 0x85)
            i += 1
            continue

        # ── Regular paragraph ────────────────────────────────────
        stripped = line.strip()
        # skip pure horizontal rule made of = or *
        if re.match(r'^[=*]{3,}$', stripped):
            i += 1
            continue

        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(4)
        # Remove markdown links
        stripped = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', stripped)
        apply_inline(p, stripped)
        i += 1

    # flush any remaining table
    if table_lines:
        parse_table(doc, table_lines)

    doc.save(out_path)
    print(f"[OK] Saved: {out_path}")
    print(f"     Lines processed: {len(lines)}")


if __name__ == "__main__":
    print("Converting BlueLine_LLD.md to BlueLine_LLD.docx ...")
    convert_md_to_docx(INPUT_FILE, OUTPUT_FILE)
